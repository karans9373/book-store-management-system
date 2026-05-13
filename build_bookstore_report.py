from __future__ import annotations

import math
import sqlite3
from collections import Counter
from datetime import date
from pathlib import Path
from typing import Iterable

import fitz
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = BASE_DIR / "deliverables"
ASSET_DIR = OUTPUT_DIR / "report_assets"
APP_FILE = BASE_DIR / "app.py"
DB_FILE = BASE_DIR / "bookverse.db"
DOCX_FILE = OUTPUT_DIR / "Book_Store_Management_System_Report.docx"

SCREENSHOT_DIR = Path(r"C:\Users\mamta\OneDrive\Pictures\Screenshots")
SCREENSHOTS = [
    ("Screenshot 2026-05-11 210009.png", "Storefront catalog hero with live inventory filters",
     "The main bookstore page communicates the product proposition immediately: a large searchable catalog, live stock metrics, and professional navigation for both customer and admin journeys. The screenshot demonstrates that the system is not limited to backend record management; it also functions as a polished customer-facing commerce interface."),
    ("Screenshot 2026-05-11 210044.png", "Community module with challenges, discussion rooms, and leaderboard",
     "The community screen shows how the application extends beyond transactional selling into reader retention and engagement. Reading challenges, discussion prompts, and leaderboard cards create a lightweight social layer that can improve repeat visits and enrich the overall platform value proposition."),
    ("Screenshot 2026-05-11 210104.png", "Admin analytics dashboard with operational KPIs",
     "This dashboard provides at-a-glance visibility into revenue, orders, inventory units, visitor traffic, and stock health. The layout supports managerial monitoring by surfacing daily snapshot indicators as well as cumulative business performance figures."),
    ("Screenshot 2026-05-11 210117.png", "Admin low-stock alerts and recent order queue",
     "The low-stock panel and recent order panel support day-to-day operational decision making. By viewing stock depletion and the latest customer transactions in the same screen, an administrator can prioritize replenishment and order handling without switching between disconnected systems."),
    ("Screenshot 2026-05-11 210132.png", "Genres page summarizing title counts, ratings, and price bands",
     "The genres page transforms raw inventory into a browsable category experience. Each genre card aggregates count, average rating, and entry price, helping users narrow their choice set while also supporting merchandising strategies for catalog exploration."),
    ("Screenshot 2026-05-11 210146.png", "Authors directory with inventory counts and profile cards",
     "Author cards show how the system groups multiple titles under a recognizable literary identity. The presentation is useful both for readers who follow specific authors and for administrators who want an overview of author concentration inside the inventory."),
    ("Screenshot 2026-05-11 210159.png", "Book clubs page with scheduled reading rooms",
     "The reading room module reflects the social-book-community objective of the project. Scheduled clubs such as genre-based groups or exam-preparation circles improve user retention by turning the platform into an interactive reading environment rather than a static store."),
    ("Screenshot 2026-05-11 210212.png", "Inventory management page with title totals and latest stock slice",
     "This screen represents the administrative inventory module. It combines aggregate counts with tabular stock data, enabling new-book entry, stock review, and deletion actions through a compact management interface."),
    ("Screenshot 2026-05-11 210236.png", "Admin order management screen showing shipment address and purchased items",
     "The order management interface captures the complete order context: customer identity, shipment address, order status, payment state, book covers, book names, and line prices. Such consolidated visibility is essential for fulfillment and support operations."),
    ("Screenshot 2026-05-11 210249.png", "Printable order detail layout",
     "A separate printable order screen improves the professional completeness of the system. This feature is valuable for documentation, dispatch verification, and producing physical records without redesigning the content specifically for print."),
    ("Screenshot 2026-05-11 210310.png", "Customer order tracking screen with delivery timeline",
     "The tracking interface mirrors familiar commercial patterns used by major e-commerce platforms. Its step-based shipment timeline clarifies the state transition from confirmation to delivery and improves customer trust after checkout."),
    ("Screenshot 2026-05-11 210412.png", "Store page showing live book cards, cover images, and preview actions",
     "Book cards integrate commerce and reading features in one compact unit. Each card contains real cover art, metadata, ratings, preview controls, wishlist action, and add-to-cart action, making the browsing experience closer to a commercial bookstore than to a simple student CRUD form."),
    ("Screenshot 2026-05-11 210444.png", "Reader interface with font controls and immersive page layout",
     "The reader page demonstrates the digital-reading component of the application. Users can read sample content, switch display modes, adjust font size, add bookmarks, and record notes, which aligns the project with hybrid bookstore-reader platforms rather than traditional inventory systems."),
    ("Screenshot 2026-05-11 210511.png", "Book detail page with stock, price, summary, and preview entry point",
     "This detail screen functions as a conversion-oriented product page. It combines bibliographic information, rating signals, availability, and summary content while also providing immediate access to preview and purchase actions."),
    ("Screenshot 2026-05-11 210601.png", "Homepage recommendation hooks: AI Book Match Quiz and BookSoul Match",
     "The AI recommendation area differentiates the project by introducing a mood-based quiz and a reader-compatibility concept. These features are useful from a project-evaluation standpoint because they show novelty, personalization, and a stronger product mindset."),
    ("Screenshot 2026-05-11 210628.png", "Homepage arrivals and best-seller showcase",
     "This merchandising section highlights new arrivals and popular books using a premium card-based visual style. It supports discovery, upselling, and the perception of an active, curated catalog."),
    ("Screenshot 2026-05-11 210728.png", "AI store assistant returning purchasable results",
     "The assistant interface shows the conversational recommendation flow. A user can type a subject such as science or mystery and receive matching available titles with short descriptions and buying actions, which improves usability for non-technical users and demonstrates intelligent search assistance."),
]


def ensure_dirs() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)


def connect_db() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_FILE)
    conn.row_factory = sqlite3.Row
    return conn


def pil_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = []
    if bold:
        candidates.extend(
            [
                r"C:\Windows\Fonts\timesbd.ttf",
                r"C:\Windows\Fonts\georgiab.ttf",
                r"C:\Windows\Fonts\cambria.ttc",
            ]
        )
    else:
        candidates.extend(
            [
                r"C:\Windows\Fonts\times.ttf",
                r"C:\Windows\Fonts\georgia.ttf",
                r"C:\Windows\Fonts\cambria.ttc",
            ]
        )
    for path in candidates:
        try:
            return ImageFont.truetype(path, size=size)
        except OSError:
            continue
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        trial = word if not current else current + " " + word
        if draw.textbbox((0, 0), trial, font=font)[2] <= width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_rounded_box(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], fill: tuple[int, int, int], outline: tuple[int, int, int], radius: int = 22, width: int = 3) -> None:
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: tuple[int, int, int], width: int = 4) -> None:
    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    arrow_len = 16
    left = (
        int(end[0] - arrow_len * math.cos(angle - math.pi / 6)),
        int(end[1] - arrow_len * math.sin(angle - math.pi / 6)),
    )
    right = (
        int(end[0] - arrow_len * math.cos(angle + math.pi / 6)),
        int(end[1] - arrow_len * math.sin(angle + math.pi / 6)),
    )
    draw.polygon([end, left, right], fill=color)


def write_centered(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font: ImageFont.ImageFont, fill: tuple[int, int, int]) -> None:
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, font, x2 - x1 - 24)
    heights = [draw.textbbox((0, 0), line, font=font)[3] for line in lines]
    total_h = sum(heights) + max(0, len(lines) - 1) * 8
    y = y1 + (y2 - y1 - total_h) // 2
    for line, h in zip(lines, heights):
        w = draw.textbbox((0, 0), line, font=font)[2]
        x = x1 + (x2 - x1 - w) // 2
        draw.text((x, y), line, font=font, fill=fill)
        y += h + 8


def create_diagram_canvas(title: str, subtitle: str, size: tuple[int, int] = (1600, 980)) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", size, (248, 247, 242))
    draw = ImageDraw.Draw(image)
    title_font = pil_font(44, bold=True)
    sub_font = pil_font(22)
    draw.text((70, 42), title, font=title_font, fill=(24, 53, 66))
    draw.text((72, 108), subtitle, font=sub_font, fill=(90, 106, 115))
    draw.line((70, 150, size[0] - 70, 150), fill=(202, 170, 88), width=4)
    return image, draw


def create_use_case_diagram() -> Path:
    path = ASSET_DIR / "use_case_diagram.png"
    image, draw = create_diagram_canvas("Use Case Diagram", "Primary interactions between customer, administrator, and the BOOKVERSE AI platform")
    title_font = pil_font(30, bold=True)
    text_font = pil_font(22)
    actor_font = pil_font(28, bold=True)
    system_box = (430, 190, 1180, 880)
    draw_rounded_box(draw, system_box, (255, 255, 255), (78, 119, 129), radius=28, width=4)
    draw.text((690, 208), "BOOKVERSE AI System", font=title_font, fill=(20, 50, 61))

    use_cases = [
        ((610, 290, 1010, 350), "Browse catalog"),
        ((610, 375, 1010, 435), "Search and filter books"),
        ((610, 460, 1010, 520), "Read 10-page preview"),
        ((610, 545, 1010, 605), "Add to cart / checkout"),
        ((610, 630, 1010, 690), "Track order"),
        ((610, 715, 1010, 775), "Use AI recommendation features"),
        ((610, 800, 1010, 860), "Manage inventory / analytics"),
    ]
    for box, label in use_cases:
        draw_rounded_box(draw, box, (245, 251, 252), (96, 139, 147), radius=24, width=3)
        write_centered(draw, box, label, text_font, (20, 39, 47))

    actors = {
        "Customer": (120, 360),
        "Administrator": (1320, 360),
    }
    for name, (x, y) in actors.items():
        draw.ellipse((x - 38, y - 88, x + 38, y - 12), outline=(40, 74, 89), width=4)
        draw.line((x, y - 12, x, y + 84), fill=(40, 74, 89), width=4)
        draw.line((x - 48, y + 18, x + 48, y + 18), fill=(40, 74, 89), width=4)
        draw.line((x, y + 84, x - 40, y + 140), fill=(40, 74, 89), width=4)
        draw.line((x, y + 84, x + 40, y + 140), fill=(40, 74, 89), width=4)
        tw = draw.textbbox((0, 0), name, font=actor_font)[2]
        draw.text((x - tw // 2, y + 158), name, font=actor_font, fill=(20, 50, 61))

    customer_links = [(158, 378, 610, 320), (158, 378, 610, 405), (158, 378, 610, 490), (158, 378, 610, 575), (158, 378, 610, 660), (158, 378, 610, 745)]
    for x1, y1, x2, y2 in customer_links:
        draw.line((x1, y1, x2, y2), fill=(173, 128, 46), width=3)
    admin_links = [(1282, 378, 1010, 320), (1282, 378, 1010, 405), (1282, 378, 1010, 830), (1282, 378, 1010, 660)]
    for x1, y1, x2, y2 in admin_links:
        draw.line((x1, y1, x2, y2), fill=(54, 120, 109), width=3)

    image.save(path)
    return path


def create_er_diagram() -> Path:
    path = ASSET_DIR / "er_diagram.png"
    image, draw = create_diagram_canvas("ER Diagram", "Core relational entities used by the Book Store Management System")
    head_font = pil_font(24, bold=True)
    body_font = pil_font(18)
    entities = {
        "books": ((110, 220, 470, 470), ["PK id", "title", "author", "genre", "price", "rating", "stock", "sold_count"]),
        "orders": ((620, 220, 980, 470), ["PK id", "order_number", "customer_name", "phone", "address", "total", "status", "placed_on"]),
        "order_items": ((1120, 220, 1480, 470), ["PK id", "FK order_id", "FK book_id", "title", "qty", "price", "line_total"]),
        "reviews": ((110, 570, 470, 800), ["PK id", "FK book_id", "reviewer", "rating", "comment"]),
        "wishlist": ((620, 570, 980, 800), ["PK id", "user_name", "FK book_id"]),
        "borrow_records": ((1120, 570, 1480, 800), ["PK id", "member_name", "FK book_id", "issue_date", "due_date", "returned", "fine"]),
    }
    for name, (box, fields) in entities.items():
        draw_rounded_box(draw, box, (255, 255, 255), (70, 114, 124), radius=24, width=4)
        x1, y1, x2, y2 = box
        draw.rectangle((x1, y1, x2, y1 + 56), fill=(235, 246, 247), outline=None)
        draw.text((x1 + 18, y1 + 14), name.upper(), font=head_font, fill=(19, 47, 56))
        y = y1 + 76
        for field in fields:
            draw.text((x1 + 22, y), field, font=body_font, fill=(35, 48, 54))
            y += 28
    draw_arrow(draw, (470, 345), (620, 345), (173, 128, 46))
    draw_arrow(draw, (980, 345), (1120, 345), (173, 128, 46))
    draw_arrow(draw, (290, 470), (290, 570), (54, 120, 109))
    draw_arrow(draw, (810, 570), (810, 470), (54, 120, 109))
    draw_arrow(draw, (1290, 570), (1290, 470), (54, 120, 109))
    image.save(path)
    return path


def create_dfd_context() -> Path:
    path = ASSET_DIR / "dfd_context.png"
    image, draw = create_diagram_canvas("Data Flow Diagram (Context Level)", "External entities and major information exchanges")
    head_font = pil_font(28, bold=True)
    body_font = pil_font(20)
    center = (590, 260, 1030, 720)
    draw_rounded_box(draw, center, (255, 255, 255), (73, 115, 124), radius=40, width=5)
    write_centered(draw, center, "BOOKVERSE AI\nBook Store Management System", head_font, (25, 52, 63))
    entities = {
        "Customer": (120, 280, 360, 420),
        "Administrator": (1250, 280, 1490, 420),
        "Public Metadata APIs": (120, 560, 420, 700),
        "Courier / Tracking View": (1200, 560, 1490, 700),
    }
    for name, box in entities.items():
        draw_rounded_box(draw, box, (246, 249, 250), (152, 167, 173), radius=24, width=3)
        write_centered(draw, box, name, body_font, (33, 45, 52))
    draw_arrow(draw, (360, 350), (590, 350), (173, 128, 46))
    draw_arrow(draw, (590, 405), (360, 405), (173, 128, 46))
    draw_arrow(draw, (1030, 350), (1250, 350), (54, 120, 109))
    draw_arrow(draw, (1250, 405), (1030, 405), (54, 120, 109))
    draw_arrow(draw, (420, 630), (690, 630), (173, 128, 46))
    draw_arrow(draw, (1130, 630), (1490, 630), (54, 120, 109))
    image.save(path)
    return path


def create_architecture_diagram() -> Path:
    path = ASSET_DIR / "system_architecture.png"
    image, draw = create_diagram_canvas("System Architecture", "Layered architecture of presentation, application logic, persistence, and external integrations")
    head_font = pil_font(26, bold=True)
    body_font = pil_font(20)
    layers = [
        ((180, 220, 1420, 330), "Presentation Layer", "HTML templates, CSS styling, JavaScript microinteractions, search UI, reader UI, and admin panels"),
        ((180, 390, 1420, 500), "Application Layer", "Flask routes, session management, cart operations, AI assistant endpoints, analytics computation, and order processing"),
        ((180, 560, 1420, 670), "Data Layer", "SQLite tables for books, orders, order_items, reviews, borrow_records, wishlist, clubs, and visitor logs"),
        ((180, 730, 1420, 840), "External Services", "Open Library metadata, Project Gutenberg / Gutendex previews, browser-based screenshots, and deployment runtime"),
    ]
    for idx, (box, title, body) in enumerate(layers):
        fill = [(241, 248, 249), (248, 246, 240), (243, 248, 242), (246, 244, 249)][idx]
        outline = [(63, 121, 114), (188, 141, 57), (94, 146, 104), (120, 108, 168)][idx]
        draw_rounded_box(draw, box, fill, outline, radius=28, width=4)
        draw.text((box[0] + 24, box[1] + 18), title, font=head_font, fill=(19, 47, 56))
        lines = wrap_text(draw, body, body_font, box[2] - box[0] - 50)
        y = box[1] + 58
        for line in lines:
            draw.text((box[0] + 24, y), line, font=body_font, fill=(36, 48, 54))
            y += 32
    draw_arrow(draw, (800, 330), (800, 390), (95, 128, 136))
    draw_arrow(draw, (800, 500), (800, 560), (95, 128, 136))
    draw_arrow(draw, (800, 670), (800, 730), (95, 128, 136))
    image.save(path)
    return path


def create_checkout_flow() -> Path:
    path = ASSET_DIR / "checkout_flowchart.png"
    image, draw = create_diagram_canvas("Checkout and Order Tracking Flowchart", "Process sequence from book discovery to order confirmation and tracking")
    font = pil_font(22)
    bold = pil_font(24, bold=True)
    steps = [
        ("Browse or search books", 280),
        ("Open detail page / preview", 390),
        ("Add to cart", 500),
        ("Enter customer details", 610),
        ("Pay successfully (demo)", 720),
        ("Generate order and tracking ID", 830),
    ]
    for label, y in steps:
        box = (550, y, 1050, y + 68)
        draw_rounded_box(draw, box, (255, 255, 255), (88, 127, 136), radius=26, width=3)
        write_centered(draw, box, label, font, (28, 46, 54))
    for _, y in steps[:-1]:
        draw_arrow(draw, (800, y + 68), (800, y + 110), (173, 128, 46))
    decision = (180, 610, 460, 680)
    draw.polygon([(320, 560), (440, 645), (320, 730), (200, 645)], fill=(247, 244, 237), outline=(188, 141, 57))
    write_centered(draw, (210, 590, 430, 700), "Logged in?", bold, (73, 57, 29))
    draw_arrow(draw, (460, 645), (550, 644), (54, 120, 109))
    draw.text((460, 610), "Yes", font=font, fill=(54, 120, 109))
    draw_arrow(draw, (320, 560), (320, 460), (173, 128, 46))
    draw.text((335, 500), "No -> Sign up / login", font=font, fill=(173, 128, 46))
    image.save(path)
    return path


def create_dfd_level1() -> Path:
    path = ASSET_DIR / "dfd_level1.png"
    image, draw = create_diagram_canvas("Data Flow Diagram (Level 1)", "Decomposition of the internal processing functions")
    font = pil_font(22)
    bold = pil_font(25, bold=True)
    processes = {
        "1.0 Catalog\nManagement": (160, 250, 450, 380),
        "2.0 Reader\nEngine": (520, 250, 810, 380),
        "3.0 Order\nProcessing": (880, 250, 1170, 380),
        "4.0 Admin\nAnalytics": (1240, 250, 1530, 380),
    }
    data_stores = {
        "D1 Books": (200, 620, 430, 720),
        "D2 Orders": (620, 620, 850, 720),
        "D3 Reader Cache": (1040, 620, 1270, 720),
        "D4 Visitor Log": (1360, 620, 1590, 720),
    }
    for label, box in processes.items():
        draw_rounded_box(draw, box, (255, 255, 255), (76, 117, 126), radius=26, width=4)
        write_centered(draw, box, label, bold, (23, 49, 58))
    for label, box in data_stores.items():
        draw.rectangle(box, fill=(245, 248, 249), outline=(118, 142, 149), width=3)
        write_centered(draw, box, label, font, (36, 50, 56))
    connectors = [
        ((305, 380), (305, 620), (173, 128, 46)),
        ((665, 380), (665, 620), (54, 120, 109)),
        ((1025, 380), (1155, 620), (173, 128, 46)),
        ((1385, 380), (1475, 620), (54, 120, 109)),
        ((450, 315), (520, 315), (95, 128, 136)),
        ((810, 315), (880, 315), (95, 128, 136)),
        ((1170, 315), (1240, 315), (95, 128, 136)),
    ]
    for start, end, color in connectors:
        draw_arrow(draw, start, end, color)
    image.save(path)
    return path


def create_bar_chart(title: str, subtitle: str, values: list[tuple[str, int]], filename: str, accent: tuple[int, int, int]) -> Path:
    path = ASSET_DIR / filename
    image, draw = create_diagram_canvas(title, subtitle, size=(1600, 900))
    label_font = pil_font(20)
    bold = pil_font(22, bold=True)
    chart_left, chart_top, chart_right, chart_bottom = 160, 240, 1480, 760
    draw.line((chart_left, chart_bottom, chart_right, chart_bottom), fill=(120, 130, 138), width=3)
    draw.line((chart_left, chart_top, chart_left, chart_bottom), fill=(120, 130, 138), width=3)
    max_val = max(v for _, v in values) if values else 1
    bar_gap = 40
    bar_width = max(50, (chart_right - chart_left - bar_gap * (len(values) + 1)) // max(len(values), 1))
    for idx, (label, value) in enumerate(values):
        x1 = chart_left + bar_gap + idx * (bar_width + bar_gap)
        x2 = x1 + bar_width
        height = int((value / max_val) * (chart_bottom - chart_top - 50))
        y1 = chart_bottom - height
        draw.rounded_rectangle((x1, y1, x2, chart_bottom), radius=18, fill=accent)
        draw.text((x1, y1 - 30), str(value), font=bold, fill=(31, 43, 51))
        lines = wrap_text(draw, label, label_font, bar_width + 30)
        y = chart_bottom + 18
        for line in lines:
            draw.text((x1 - 6, y), line, font=label_font, fill=(31, 43, 51))
            y += 24
    image.save(path)
    return path


def create_assets() -> dict[str, Path]:
    conn = connect_db()
    genre_rows = conn.execute("SELECT genre, COUNT(*) AS total FROM books GROUP BY genre ORDER BY total DESC LIMIT 8").fetchall()
    status_rows = conn.execute("SELECT status, COUNT(*) AS total FROM orders GROUP BY status ORDER BY total DESC").fetchall()
    author_rows = conn.execute("SELECT author, COUNT(*) AS total FROM books GROUP BY author ORDER BY total DESC LIMIT 6").fetchall()
    conn.close()
    return {
        "use_case": create_use_case_diagram(),
        "er": create_er_diagram(),
        "dfd_context": create_dfd_context(),
        "architecture": create_architecture_diagram(),
        "checkout_flow": create_checkout_flow(),
        "dfd_level1": create_dfd_level1(),
        "genre_chart": create_bar_chart(
            "Genre Distribution in Live Catalog",
            "Top genres imported into the BOOKVERSE AI inventory",
            [(row["genre"], row["total"]) for row in genre_rows],
            "genre_distribution.png",
            (54, 120, 109),
        ),
        "status_chart": create_bar_chart(
            "Order Status Distribution",
            "Seed and live order states recorded in the order management module",
            [(row["status"], row["total"]) for row in status_rows],
            "order_status_distribution.png",
            (188, 141, 57),
        ),
        "author_chart": create_bar_chart(
            "Author Representation",
            "Top authors by number of titles currently present in inventory",
            [(row["author"], row["total"]) for row in author_rows],
            "author_representation.png",
            (76, 117, 176),
        ),
    }


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Table of contents will be generated during document update."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def configure_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.35)
    sec.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in [
        ("Title", 24, RGBColor(22, 47, 60)),
        ("Heading 1", 17, RGBColor(20, 60, 73)),
        ("Heading 2", 14, RGBColor(26, 55, 70)),
        ("Heading 3", 12, RGBColor(39, 64, 77)),
    ]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

    if "CaptionCustom" not in styles:
        cap = styles.add_style("CaptionCustom", WD_STYLE_TYPE.PARAGRAPH)
        cap.font.name = "Times New Roman"
        cap._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        cap.font.size = Pt(10)
        cap.font.italic = True
        cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(4)
        cap.paragraph_format.space_after = Pt(8)

    if "CodeBlock" not in styles:
        code = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        code.font.name = "Consolas"
        code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        code.font.size = Pt(9.5)
        code.paragraph_format.line_spacing = 1.15
        code.paragraph_format.space_before = Pt(0)
        code.paragraph_format.space_after = Pt(0)

    if "SmallText" not in styles:
        sm = styles.add_style("SmallText", WD_STYLE_TYPE.PARAGRAPH)
        sm.font.name = "Times New Roman"
        sm._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        sm.font.size = Pt(10)
        sm.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        sm.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Page ")
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        add_page_number(p)


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    run = p.add_run("BOOK STORE MANAGEMENT SYSTEM")
    run.font.name = "Times New Roman"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(20, 50, 61)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("A Comprehensive Project Report on the Design and Implementation of BOOKVERSE AI")
    run.font.name = "Times New Roman"
    run.font.size = Pt(15)
    run.italic = True

    doc.add_paragraph("")
    center_lines = [
        "Submitted in partial fulfillment of the requirements for the award of the degree of",
        "Bachelor of Computer Applications / B.Sc. (Computer Science) / B.Tech. (Computer Science)",
        "Academic Session: 2025-2026",
        "",
        "Submitted By",
        "Karan S. (Sample Student Presentation Identity)",
        "Roll No.: ____________________",
        "",
        "Under the Guidance of",
        "Project Guide / Faculty Mentor",
        "",
        "Department of Computer Applications / Computer Science",
        "____________________________ College / University",
        f"Submission Date: {date.today():%d %B %Y}",
    ]
    for line in center_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.name = "Times New Roman"
        run.font.size = Pt(13 if line and "Submitted By" not in line and "Under the Guidance" not in line else 14)
        if line in {"Submitted By", "Under the Guidance of"}:
            run.bold = True


def add_formal_page(doc: Document, title: str, paragraphs: list[str]) -> None:
    add_page_break(doc)
    p = doc.add_paragraph(title, style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for text in paragraphs:
        para = doc.add_paragraph(style="Normal")
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run(text)


def add_toc_page(doc: Document, toc_rows: list[tuple[str, str]] | None = None) -> None:
    add_page_break(doc)
    doc.add_paragraph("Table of Contents", style="Heading 1").alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not toc_rows:
        p = doc.add_paragraph()
        add_toc(p)
        p = doc.add_paragraph("The following edition includes chapter and section references after the final layout pass.", style="SmallText")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return

    for title, page in toc_rows:
        p = doc.add_paragraph(style="SmallText")
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), alignment=2, leader=1)
        if title.startswith("    "):
            p.paragraph_format.left_indent = Inches(0.35)
            run = p.add_run(title.strip())
        else:
            run = p.add_run(title)
            run.bold = True
        p.add_run("\t" + (page or ""))
    doc.add_paragraph("Note: Chapters begin on fresh pages; section references were aligned against the rendered PDF used for final formatting review.", style="SmallText")


def add_abbreviation_page(doc: Document) -> None:
    add_page_break(doc)
    doc.add_paragraph("List of Symbols, Abbreviations and Nomenclature", style="Heading 1")
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, text in enumerate(["Abbreviation", "Full Form", "Meaning in Project Context"]):
        hdr[idx].text = text
        set_cell_shading(hdr[idx], "EAF5F7")
    rows = [
        ("AI", "Artificial Intelligence", "Used for mood-based suggestion, conversational recommendations, and decision support."),
        ("API", "Application Programming Interface", "Interface used to fetch public metadata and previews from external data sources."),
        ("BCA", "Bachelor of Computer Applications", "One of the academic programs for which the report structure is suitable."),
        ("CRUD", "Create, Read, Update, Delete", "Basic record operations often found in simple systems; this project extends far beyond CRUD."),
        ("DFD", "Data Flow Diagram", "Represents the movement of information through system processes and stores."),
        ("ER", "Entity Relationship", "Shows database entities such as books, orders, and order_items with their relationships."),
        ("GUI", "Graphical User Interface", "Visual interface presented to customer and administrator users."),
        ("GUI", "Graphical User Interface", "Visual interface presented to customer and administrator users."),
        ("HTML", "HyperText Markup Language", "Used for template rendering in the frontend layer."),
        ("HTTP", "HyperText Transfer Protocol", "Protocol used for route requests and responses in the web application."),
        ("IDE", "Integrated Development Environment", "Development tool used for coding and debugging."),
        ("JSON", "JavaScript Object Notation", "Format used for API payloads, cached reader pages, and route responses."),
        ("KPI", "Key Performance Indicator", "Analytics values such as revenue, orders, visitors, and inventory units."),
        ("PDF", "Portable Document Format", "Used for the final printable version of the report and order exports."),
        ("RDBMS", "Relational Database Management System", "SQLite serves as the structured persistence layer for the project."),
        ("SQL", "Structured Query Language", "Used for schema design, data retrieval, stock updates, and order persistence."),
        ("SVG", "Scalable Vector Graphics", "Used for generated placeholder covers and lightweight visual assets."),
        ("UML", "Unified Modeling Language", "Diagram standard used for use case, activity, and structure modelling."),
        ("URL", "Uniform Resource Locator", "Used for references to routes, metadata endpoints, and cover/image sources."),
        ("UX", "User Experience", "Relates to navigation quality, readability, and interaction smoothness."),
    ]
    for item in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(item):
            cells[idx].text = value
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_chapter_start(doc: Document, chapter_no: int, title: str) -> None:
    add_page_break(doc)
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(f"Chapter {chapter_no}: {title}")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(20, 60, 73)


def add_section(doc: Document, heading: str, paragraphs: Iterable[str], bullet_points: list[str] | None = None) -> None:
    doc.add_paragraph(heading, style="Heading 2")
    for text in paragraphs:
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(text)
    if bullet_points:
        for point in bullet_points:
            p = doc.add_paragraph(style="Normal")
            p.style = "List Bullet"
            p.add_run(point)


def add_table(doc: Document, title: str, headers: list[str], rows: list[tuple[str, ...] | list[str]]) -> None:
    doc.add_paragraph(title, style="Heading 3")
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, head in enumerate(headers):
        hdr_cells[i].text = head
        set_cell_shading(hdr_cells[i], "EEF5F7")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph("")


def add_figure(doc: Document, image_path: Path, caption: str, explanation: str, width: float = 6.2) -> None:
    doc.add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph(f"Figure: {caption}", style="CaptionCustom")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para = doc.add_paragraph(style="Normal")
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.add_run(explanation)


def add_code_block(doc: Document, title: str, code: str, explanation: str) -> None:
    doc.add_paragraph(title, style="Heading 2")
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F7F7")
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.left_indent = Inches(0.15)
    para.paragraph_format.right_indent = Inches(0.15)
    for idx, line in enumerate(code.splitlines()):
        if idx:
            para.add_run("\n")
        run = para.add_run(line.rstrip())
        run.font.name = "Consolas"
        run.font.size = Pt(9)
    doc.add_paragraph(explanation, style="Normal")


def fetch_live_stats() -> dict[str, int | str]:
    conn = connect_db()
    stats = {
        "books": conn.execute("SELECT COUNT(*) FROM books").fetchone()[0],
        "inventory_units": conn.execute("SELECT COALESCE(SUM(stock), 0) FROM books").fetchone()[0],
        "orders": conn.execute("SELECT COUNT(*) FROM orders").fetchone()[0],
        "revenue": conn.execute("SELECT COALESCE(SUM(total), 0) FROM orders").fetchone()[0],
        "visitors": conn.execute("SELECT COUNT(*) FROM visitor_log").fetchone()[0],
        "reviews": conn.execute("SELECT COUNT(*) FROM reviews").fetchone()[0],
        "clubs": conn.execute("SELECT COUNT(*) FROM clubs").fetchone()[0],
        "preview_books": conn.execute("SELECT COUNT(*) FROM books WHERE text_url <> ''").fetchone()[0],
    }
    stats["top_genres"] = ", ".join(
        row["genre"]
        for row in conn.execute("SELECT genre, COUNT(*) AS total FROM books GROUP BY genre ORDER BY total DESC LIMIT 5").fetchall()
    )
    conn.close()
    return stats


def extract_snippet(start_marker: str, end_marker: str | None = None, max_lines: int = 42) -> str:
    text = APP_FILE.read_text(encoding="utf-8")
    start = text.find(start_marker)
    if start == -1:
        return start_marker
    if end_marker:
        end = text.find(end_marker, start + len(start_marker))
        snippet = text[start:end if end != -1 else None]
    else:
        snippet = text[start:]
    lines = snippet.strip().splitlines()[:max_lines]
    return "\n".join(lines)


def make_paragraphs(section_name: str, themes: list[str], chapter_context: str, module_detail: str) -> list[str]:
    theme_line = ", ".join(themes[:-1]) + (", and " + themes[-1] if len(themes) > 1 else themes[0])
    p1 = (
        f"The section titled '{section_name}' is central to the academic understanding of the proposed Book Store Management System because it explains how a modern digital bookstore must operate as more than a data-entry utility. "
        f"In this project, the BOOKVERSE AI platform combines catalog presentation, order management, preview reading, and administrative supervision into one integrated application. "
        f"From a systems perspective, {section_name.lower()} is approached through the lenses of {theme_line}. "
        f"This discussion is intentionally grounded in the implemented Flask application so that the report reflects an actual engineering artifact rather than a hypothetical classroom concept."
    )
    p2 = (
        f"In conventional book-selling or library-style workflows, many of these activities are fragmented across notebooks, spreadsheets, isolated web forms, or separate applications. "
        f"Such fragmentation increases information latency, introduces stock inconsistencies, and weakens the user experience because the reader, buyer, and administrator all perceive different versions of the same system state. "
        f"The present project addresses this limitation by maintaining a consistent operational model where catalog availability, reader previews, cart actions, tracking records, and managerial analytics remain traceable inside one coherent solution. "
        f"This is particularly relevant in academic project evaluation because it demonstrates integration discipline rather than feature accumulation."
    )
    p3 = (
        f"Within BOOKVERSE AI, {module_detail} is implemented using a combination of Flask route handlers, server-rendered templates, SQLite-backed persistence, and structured session-driven user flows. "
        f"These technical choices were selected for clarity, portability, and pedagogical value. "
        f"The architecture remains lightweight enough for student deployment while still supporting realistic e-commerce and management behaviors such as stock decrement after purchase, order number generation, printable invoices, tracking milestones, recommendation prompts, and administrative monitoring. "
        f"Therefore, the section does not merely describe theory; it demonstrates how theory is converted into practical software behavior."
    )
    p4 = (
        f"From an academic reporting standpoint, the significance of {section_name.lower()} also lies in its contribution to maintainability and future extensibility. "
        f"When project stakeholders can observe clearly defined data responsibilities, interface expectations, and operational outcomes, the software becomes easier to test, document, refine, and migrate. "
        f"In the context of the chosen project, the {chapter_context} viewpoint helps explain why the solution can support both present requirements and future growth such as payment-gateway integration, hosted databases, recommendation refinement, and broader reader-community functions."
    )
    return [p1, p2, p3]


def add_reference_list(doc: Document, references: list[str]) -> None:
    for ref in references:
        p = doc.add_paragraph(style="Normal")
        p.add_run(ref)


def extract_heading_pages(pdf_path: Path, headings: list[str]) -> dict[str, int]:
    if not pdf_path.exists():
        return {}
    doc = fitz.open(pdf_path)
    found: dict[str, int] = {}
    for idx, page in enumerate(doc, start=1):
        text = page.get_text("text")
        normalized = " ".join(text.split())
        for heading in headings:
            if heading in found:
                continue
            if heading in normalized:
                found[heading] = idx
    doc.close()
    return found


def build_report() -> None:
    ensure_dirs()
    assets = create_assets()
    stats = fetch_live_stats()
    pdf_path = OUTPUT_DIR / "rendered_report" / "Book_Store_Management_System_Report.pdf"

    doc = Document()
    configure_document(doc)
    add_title_page(doc)

    add_formal_page(
        doc,
        "Certificate",
        [
            "This is to certify that the project report entitled 'Book Store Management System' submitted by the student in partial fulfillment of the requirements for the award of the degree in Computer Applications / Computer Science is a bona fide record of work carried out under our supervision during the academic session 2025-2026.",
            "The work presented in this report has been examined for academic relevance, implementation effort, and technical coherence. To the best of our knowledge, the project demonstrates an integrated understanding of web application development, database management, user experience design, and administrative reporting in the domain of digital book commerce and management.",
            "Project Guide Signature: ____________________    Head of Department: ____________________    Principal: ____________________",
        ],
    )
    add_formal_page(
        doc,
        "Declaration",
        [
            "I hereby declare that the project report entitled 'Book Store Management System' is an original piece of work carried out by me under the guidance of the concerned faculty member. The matter embodied in this report has not been submitted elsewhere for the award of any degree, diploma, certificate, or academic credit.",
            "I further declare that appropriate acknowledgement has been provided wherever external sources, public APIs, framework documentation, and reference materials have informed the development or explanation of the project. The software implementation, screenshots, database structure, and report compilation represent my own academic submission for evaluation.",
            "Student Signature: ____________________",
        ],
    )
    add_formal_page(
        doc,
        "Acknowledgement",
        [
            "The completion of this project report and the corresponding web application has been possible through the support of faculty mentors, institutional resources, and the availability of open technical documentation. I express my sincere gratitude to my project guide for providing direction on problem formulation, application structuring, and technical documentation standards.",
            "I also acknowledge the role of the departmental academic environment, which encouraged the selection of a project topic that merges e-commerce, reading technology, and management analytics into one practical system. The opportunity to work on a real full-stack implementation significantly strengthened my understanding of software engineering beyond isolated laboratory assignments.",
            "Finally, I acknowledge the contribution of open reference ecosystems such as Flask documentation, SQLite guidance, Open Library metadata resources, and Project Gutenberg / Gutendex preview resources, which informed the implementation and made it possible to design a richer and more realistic Book Store Management System.",
        ],
    )
    add_formal_page(
        doc,
        "Abstract",
        [
            f"The Book Store Management System presented in this report is a full-stack web application implemented using Python Flask, HTML5, CSS3, JavaScript, and SQLite. The system has been designed as a hybrid platform that serves both customer-facing and administrative requirements. In its current implemented form, the application manages a live catalog of approximately {stats['books']} books and {stats['inventory_units']} available inventory units, while also supporting reader previews, wishlist handling, cart and checkout workflows, order tracking, administrative inventory updates, and analytics dashboards.",
            "Unlike a conventional student CRUD project that focuses narrowly on add, update, delete, and search operations, this system adopts a product-oriented design philosophy. It combines commercial bookstore features with management operations, recommendation flows, and social reading functions. The implemented modules include a premium homepage, searchable and filterable catalog, book detail pages, in-browser sample reader, AI store assistant, mood-based recommendation hooks, wishlist and cart management, order processing, printable order details, track-order flow, community space, author and genre directories, and administrator-oriented inventory and analytics interfaces.",
            "The objective of the project is to demonstrate how a single digital platform can manage catalog operations, improve customer convenience, and support administrative decision making in the context of online book retail and reading engagement. The report discusses system analysis, feasibility, design diagrams, database design, implementation logic, testing approach, screenshots, source-code excerpts, and results. The final outcome is a submission-ready web-based Book Store Management System that is academically valid, technically coherent, and suitable for demonstration in a university evaluation setting.",
        ],
    )
    chapter_specs = [
        (1, "Introduction", [
            ("Background of the Study", ["digital commerce transformation", "reader expectations", "inventory visibility", "interactive interfaces"], "introductory analysis", "the combined customer and admin workflows"),
            ("Need for a Web-Based Book Store Management System", ["catalog scale", "transaction accuracy", "reader convenience", "administrative control"], "problem framing", "the search, cart, and order-management pipeline"),
            ("Overview of BOOKVERSE AI", ["premium user interface", "reader preview capability", "analytics dashboard", "recommendation hooks"], "solution overview", "the end-to-end BOOKVERSE AI experience"),
            ("Organization of the Report", ["chapter sequencing", "technical documentation", "implementation evidence", "evaluation readiness"], "report navigation", "the structure of this academic submission"),
        ]),
        (2, "Literature Survey", [
            ("Traditional Bookstore and Library Systems", ["manual operations", "book issue registers", "paper dependency", "limited discoverability"], "comparative literature review", "the transition from manual records to unified digital workflows"),
            ("Evolution of Online Book Selling Platforms", ["product catalogs", "search and recommendation", "payment convenience", "customer retention"], "market-oriented review", "commercial bookstore interaction patterns"),
            ("Digital Reading and Preview Platforms", ["sample reading", "font adaptation", "session continuity", "reader engagement"], "reader-experience review", "the implemented preview-reader module"),
            ("Academic Gap Addressed by the Present Project", ["integration gap", "project novelty", "student-level implementation", "multi-role operation"], "research positioning", "the integration of store, reader, and admin functions"),
        ]),
        (3, "Problem Statement", [
            ("Operational Problems in Fragmented Book Management", ["data inconsistency", "stock mismatch", "delayed reporting", "poor coordination"], "problem articulation", "the replacement of fragmented data handling"),
            ("User Experience Problems in Basic College Projects", ["plain CRUD layouts", "weak navigation", "non-commercial feel", "limited realism"], "design critique", "the premium storefront and management presentation"),
            ("Statement of the Core Problem", ["integration requirement", "commerce and management unity", "reader engagement", "decision support"], "project objective framing", "the full system behaviour"),
        ]),
        (4, "Objectives of the Project", [
            ("Primary Objective", ["develop an integrated web platform", "support catalog and order flows", "offer professional interfaces", "support academic evaluation"], "objective statement", "the complete platform baseline"),
            ("Functional Objectives", ["search", "filtering", "cart", "checkout", "admin dashboards"], "feature objective framing", "the implemented customer and administrator operations"),
            ("Quality Objectives", ["usability", "maintainability", "performance clarity", "future scalability"], "quality planning", "the software engineering characteristics of the system"),
        ]),
        (5, "Scope of the Project", [
            ("Scope for Customers", ["catalog browsing", "book detail exploration", "sample reading", "purchase tracking"], "scope delimitation", "the customer-facing commerce journey"),
            ("Scope for Administrators", ["inventory supervision", "order visibility", "printable outputs", "analytics"], "administrative scope", "the manager-facing control surfaces"),
            ("Project Boundaries and Assumptions", ["demo payment", "sample dataset", "public metadata sources", "local database"], "scope control", "the realistic but bounded student-project environment"),
        ]),
        (6, "System Analysis", [
            ("Existing System Analysis", ["manual bookkeeping", "limited reporting", "poor searchability", "process delay"], "analysis of current-state systems", "the motivation for a unified application"),
            ("Proposed System Analysis", ["responsive interface", "centralized data", "analytics readiness", "role-based visibility"], "analysis of the target state", "the modular BOOKVERSE AI workflow"),
            ("Input, Process, and Output Analysis", ["customer input", "server processing", "database writes", "presented outputs"], "IPO perspective", "the route-handler and template-response cycle"),
            ("User and Stakeholder Analysis", ["students", "general readers", "admins", "project evaluators"], "stakeholder analysis", "the alignment of features with user classes"),
        ]),
        (7, "Feasibility Study", [
            ("Technical Feasibility", ["Python Flask stack", "SQLite simplicity", "templating approach", "low deployment barrier"], "feasibility evaluation", "the selected implementation stack"),
            ("Economic Feasibility", ["low-cost development", "open frameworks", "reusable components", "student affordability"], "economic reasoning", "the cost-efficient project design"),
            ("Operational Feasibility", ["ease of use", "administrator adoption", "customer familiarity", "workflow continuity"], "operational reasoning", "the customer and admin interaction model"),
            ("Schedule Feasibility", ["modular development", "progressive implementation", "testing checkpoints", "report readiness"], "time-bound project planning", "the staged build strategy"),
        ]),
        (8, "Software & Hardware Requirements", [
            ("Software Requirement Analysis", ["Python runtime", "Flask framework", "HTML/CSS/JS", "SQLite"], "platform specification", "the selected software environment"),
            ("Development Tools and Libraries", ["VS Code or equivalent IDE", "browser testing", "image assets", "documentation tools"], "tooling description", "the supporting development toolkit"),
            ("Hardware Requirement Analysis", ["processor class", "RAM needs", "display resolution", "network access"], "resource planning", "the environment used for development and demonstration"),
            ("Deployment Considerations", ["local development server", "GitHub repository", "Render deployment readiness", "hosted execution"], "deployment planning", "the transition from local build to live hosting"),
        ]),
        (9, "System Design", [
            ("High-Level Design Strategy", ["modular separation", "route-oriented architecture", "data persistence", "template rendering"], "solution design", "the layered software structure"),
            ("Presentation Layer Design", ["navigation hierarchy", "premium visual style", "glassmorphism influence", "responsive layout"], "frontend design analysis", "the customer and admin page compositions"),
            ("Application Logic Design", ["session handling", "analytics calculation", "recommendation endpoints", "inventory updates"], "logic-layer design", "the Flask controller patterns"),
            ("Persistence and Data Integrity Design", ["schema normalization", "stock updates", "order snapshots", "cache storage"], "data-layer design", "the SQLite entity relationships"),
        ]),
        (10, "UML Diagrams", [
            ("Use Case Modelling", ["customer actor", "administrator actor", "book discovery", "order supervision"], "behavioural modelling", "the system interaction scope"),
            ("Activity and Process Modelling", ["checkout sequence", "login dependency", "status transitions", "reader navigation"], "workflow modelling", "the cart-to-tracking lifecycle"),
            ("System Interaction Decomposition", ["data flow", "processing modules", "storage points", "external interfaces"], "structural process mapping", "the internal process decomposition"),
        ]),
        (11, "Database Design", [
            ("Entity Identification and Relationship Design", ["books", "orders", "order_items", "supporting entities"], "database modelling", "the declared schema in init_db"),
            ("Normalization and Data Integrity", ["non-redundancy", "transaction capture", "stock consistency", "lookup simplicity"], "normalization review", "the relationship between master and transactional records"),
            ("Schema Behaviour in the Running Application", ["seed data", "analytics queries", "review storage", "borrow records"], "runtime schema analysis", "the application-state representation"),
            ("Query Strategy and Reporting Support", ["aggregate metrics", "genre counts", "status views", "top titles"], "query design", "the analytics-oriented SQL used in the system"),
        ]),
        (12, "Module Description", [
            ("Homepage and Discovery Module", ["hero presentation", "new arrivals", "genre cards", "hook features"], "module-level explanation", "the storefront landing experience"),
            ("Catalog and Book Detail Module", ["search", "filters", "metadata display", "action buttons"], "module-level explanation", "the store and detail-route behaviour"),
            ("Reader Module", ["10-page preview", "font controls", "bookmarks", "reading immersion"], "module-level explanation", "the in-browser reader experience"),
            ("AI Suggestion and Recommendation Module", ["store assistant", "mood quiz", "BookSoul Match", "contextual suggestions"], "module-level explanation", "the recommendation and assistant APIs"),
            ("Wishlist, Cart, and Checkout Module", ["cart persistence", "customer form capture", "order creation", "success feedback"], "module-level explanation", "the commerce transaction flow"),
            ("Order Tracking and Printable Output Module", ["tracking timeline", "printable invoice", "customer visibility", "admin handling"], "module-level explanation", "the post-purchase experience"),
            ("Admin Analytics and Inventory Module", ["dashboard KPIs", "stock alerts", "order inspection", "inventory editing"], "module-level explanation", "the managerial control surface"),
            ("Community and Engagement Module", ["challenges", "discussion prompts", "book clubs", "leaderboard"], "module-level explanation", "the social reading extension"),
        ]),
        (13, "Implementation", [
            ("Backend Implementation Using Flask", ["route design", "request handling", "session usage", "template rendering"], "implementation-level explanation", "the Flask application structure"),
            ("Database Initialization and Seed Logic", ["schema creation", "column backfill", "catalog seeding", "order bootstrapping"], "implementation-level explanation", "the init_db and seed operations"),
            ("Frontend Implementation and UI Behaviour", ["responsive cards", "interactive buttons", "visual hierarchy", "microinteractions"], "implementation-level explanation", "the templates and styling system"),
            ("Recommendation and Assistant Endpoint Logic", ["token matching", "payload generation", "context-sensitive responses", "availability-based ranking"], "implementation-level explanation", "the AI store assistant and related endpoints"),
            ("Checkout and Stock Mutation Logic", ["customer form validation", "order item creation", "stock decrement", "tracking ID generation"], "implementation-level explanation", "the order-placement transaction sequence"),
            ("Administrative Reporting and Print Workflow", ["dashboard metrics", "recent order rendering", "print route", "analytics views"], "implementation-level explanation", "the admin-monitoring and print support logic"),
        ]),
        (16, "Testing", [
            ("Testing Strategy", ["unit-oriented checks", "route validation", "UI walkthroughs", "data consistency review"], "quality-assurance explanation", "the combined functional verification approach"),
            ("Functional Test Case Design", ["authentication", "add-to-cart", "checkout", "track order"], "test-design explanation", "the major user-facing flows"),
            ("Administrative Test Case Design", ["inventory addition", "stock deletion", "analytics correctness", "printable order availability"], "test-design explanation", "the admin-facing flows"),
            ("Usability and Interface Testing", ["readability", "navigation clarity", "responsive behaviour", "button feedback"], "usability validation", "the visual and interaction quality checks"),
            ("Observed Defects and Resolution Approach", ["layout corrections", "reader-page behaviour", "dashboard readability", "login flow fixes"], "defect management reflection", "the iterative refinement of the final application"),
        ]),
        (17, "Results & Discussion", [
            ("Observed System Outcomes", ["live inventory", "integrated ordering", "reader previews", "analytics availability"], "results interpretation", "the measurable behaviour of the working system"),
            ("Interpretation of Analytics and Catalog Data", ["genre distribution", "author spread", "order states", "inventory movement"], "results interpretation", "the current state of the seeded application database"),
            ("Academic Evaluation of the Final Product", ["feature completeness", "professional appearance", "real-world resemblance", "integration quality"], "results interpretation", "the final submission strength"),
            ("Limitations of Current Observations", ["demo dataset", "simulated payment", "local database", "host dependency"], "discussion balance", "the boundary conditions of the current implementation"),
        ]),
        (18, "Advantages & Limitations", [
            ("Advantages of the Proposed System", ["unified workflow", "professional UI", "reader engagement", "admin visibility"], "comparative evaluation", "the strengths of the implemented system"),
            ("Limitations of the Current Version", ["local persistence", "demo payment model", "manual admin authentication", "preview availability constraints"], "comparative evaluation", "the areas that remain intentionally simplified"),
        ]),
        (19, "Future Enhancements", [
            ("Technical Enhancements", ["PostgreSQL migration", "payment gateway integration", "role persistence", "cloud deployment"], "future-work planning", "the likely evolution of the backend"),
            ("Product Enhancements", ["personalized shelves", "richer social interactions", "recommendation refinement", "mobile packaging"], "future-work planning", "the likely evolution of the user experience"),
            ("Academic and Commercial Extension Potential", ["library collaboration", "analytics exports", "subscription models", "multi-vendor expansion"], "future-work planning", "the broader applicability of the solution"),
        ]),
        (20, "Conclusion", [
            ("Conclusion of the Study", ["problem solution", "integration success", "technical learning", "product quality"], "concluding reflection", "the final implemented BOOKVERSE AI system"),
            ("Closing Academic Remark", ["submission readiness", "realistic scope", "engineering discipline", "future adaptability"], "concluding reflection", "the final project outcome"),
        ]),
    ]

    heading_targets = [f"Chapter {chapter_no}: {chapter_title}" for chapter_no, chapter_title, _ in chapter_specs]
    heading_targets.extend([
        "Chapter 14: Screenshots with Explanation",
        "Chapter 15: Source Code",
        "Chapter 21: References",
        "Chapter 22: Appendices",
    ])
    for _, _, sections in chapter_specs:
        heading_targets.extend([heading for heading, _, _, _ in sections])
    heading_targets.extend([
        "Appendix A: Sample Route Directory",
        "Appendix B: SQL Schema Notes",
        "Appendix C: User Manual Summary",
        "Appendix D: Observed Live Project Metrics",
    ])
    page_map = extract_heading_pages(pdf_path, heading_targets)
    toc_rows: list[tuple[str, str]] = []
    for chapter_no, chapter_title, sections in chapter_specs:
        chapter_key = f"Chapter {chapter_no}: {chapter_title}"
        toc_rows.append((chapter_key, str(page_map.get(chapter_key, ""))))
        for heading, _, _, _ in sections:
            toc_rows.append((f"    {heading}", str(page_map.get(heading, ""))))
    toc_rows.extend([
        ("Chapter 14: Screenshots with Explanation", str(page_map.get("Chapter 14: Screenshots with Explanation", ""))),
        ("Chapter 15: Source Code", str(page_map.get("Chapter 15: Source Code", ""))),
        ("Chapter 21: References", str(page_map.get("Chapter 21: References", ""))),
        ("Chapter 22: Appendices", str(page_map.get("Chapter 22: Appendices", ""))),
        ("    Appendix A: Sample Route Directory", str(page_map.get("Appendix A: Sample Route Directory", ""))),
        ("    Appendix B: SQL Schema Notes", str(page_map.get("Appendix B: SQL Schema Notes", ""))),
        ("    Appendix C: User Manual Summary", str(page_map.get("Appendix C: User Manual Summary", ""))),
        ("    Appendix D: Observed Live Project Metrics", str(page_map.get("Appendix D: Observed Live Project Metrics", ""))),
    ])

    add_toc_page(doc, toc_rows)
    add_abbreviation_page(doc)

    for chapter_no, chapter_title, sections in chapter_specs:
        add_chapter_start(doc, chapter_no, chapter_title)
        for heading, themes, chapter_context, module_detail in sections:
            add_section(doc, heading, make_paragraphs(heading, themes, chapter_context, module_detail))
        if chapter_no == 4:
            add_table(
                doc,
                "Table 4.1 Objectives mapped to implementation outcomes",
                ["Objective", "Implementation Evidence", "Expected Benefit"],
                [
                    ("Integrated catalog management", "Store, detail, and inventory routes", "Single source of truth for books"),
                    ("Professional user experience", "Premium homepage and card-based UI", "Improved readability and evaluator impression"),
                    ("Reader engagement", "10-page preview and AI suggestion interfaces", "Higher product distinctiveness"),
                    ("Administrative control", "Dashboard, orders, and print screen", "Operational decision support"),
                ],
            )
        if chapter_no == 8:
            add_table(
                doc,
                "Table 8.1 Software requirements",
                ["Component", "Specification", "Role in Project"],
                [
                    ("Programming Language", "Python 3.x", "Backend implementation"),
                    ("Framework", "Flask", "Routing, templating, and web-server logic"),
                    ("Database", "SQLite", "Lightweight relational persistence"),
                    ("Frontend Technologies", "HTML5, CSS3, JavaScript", "User interface implementation"),
                    ("Version Control", "Git and GitHub", "Source management and deployment readiness"),
                ],
            )
            add_table(
                doc,
                "Table 8.2 Hardware requirements",
                ["Hardware Item", "Minimum Requirement", "Recommended Requirement"],
                [
                    ("Processor", "Dual-core CPU", "Quad-core CPU"),
                    ("RAM", "4 GB", "8 GB or above"),
                    ("Storage", "2 GB free space", "5 GB free space"),
                    ("Display", "1366 x 768", "1920 x 1080"),
                    ("Network", "Basic internet connection", "Stable broadband connection"),
                ],
            )
        if chapter_no == 10:
            add_figure(doc, assets["use_case"], "Use Case Diagram of BOOKVERSE AI", "The use case model shows the functional scope of the system. Customer operations focus on discovery, previewing, wishlist/cart actions, checkout, and order tracking, whereas administrator operations focus on analytics, inventory control, and order supervision.", width=6.2)
            add_figure(doc, assets["checkout_flow"], "Checkout and order processing flowchart", "The flowchart shows the dependency of checkout on authentication, followed by customer detail capture, simulated payment, and automated generation of order and tracking identifiers. This sequence confirms that transactional logic and post-purchase visibility were implemented coherently.", width=6.1)
            add_figure(doc, assets["dfd_context"], "Context-level DFD", "The context-level diagram places BOOKVERSE AI at the center of exchanges with customers, administrators, public metadata APIs, and the delivery-tracking perspective. It clarifies system boundaries and external information movement.", width=6.15)
            add_figure(doc, assets["dfd_level1"], "Level-1 DFD for internal processing", "The level-1 DFD decomposes the platform into major processes such as catalog management, reader engine, order processing, and analytics. It also identifies how these processes rely on shared data stores such as books, orders, cached reader pages, and visitor logs.", width=6.15)
        if chapter_no == 11:
            add_figure(doc, assets["er"], "ER diagram for the implemented SQLite schema", "The ER diagram summarizes the relational structure that supports the current application. Books act as the master inventory entity, while orders, order_items, reviews, wishlist records, and borrow records capture transactional and engagement data.", width=6.1)
            add_figure(doc, assets["architecture"], "High-level system architecture", "The layered architecture representation explains the separation between presentation logic, application control, persistence, and external data sources. This separation improves maintainability and helps map academic design concepts to the final implementation.", width=6.15)
            add_table(
                doc,
                "Table 11.1 Core database tables",
                ["Table Name", "Purpose", "Key Attributes"],
                [
                    ("books", "Stores catalog records and stock", "title, author, genre, price, stock, cover_url, text_url"),
                    ("orders", "Stores customer purchase master data", "order_number, customer_name, address, total, status"),
                    ("order_items", "Stores item-level order lines", "order_id, book_id, qty, price, line_total"),
                    ("reviews", "Stores ratings and comments", "book_id, reviewer, rating, comment"),
                    ("visitor_log", "Stores page-view data for analytics", "visitor_key, path, viewed_on, viewed_at"),
                ],
            )
        if chapter_no == 17:
            add_figure(doc, assets["genre_chart"], "Genre distribution chart", "This chart reflects the current live database composition. It demonstrates that the application is not operating on a tiny hand-entered sample; instead, the catalog contains a broad distribution of genres, which strengthens the realism of search, filtering, and recommendation experiments.", width=6.0)
            add_figure(doc, assets["author_chart"], "Author representation chart", "Author concentration analysis helps explain inventory diversity. It also supports administrative understanding of whether the system is heavily skewed toward a few authors or broadly representative of a larger textual collection.", width=6.0)
            add_figure(doc, assets["status_chart"], "Order status distribution chart", "The order-status chart illustrates how transactional states are represented inside the dataset and user interfaces. Such classification is important for rendering tracking stages and for maintaining clarity in administrative monitoring.", width=6.0)

    add_chapter_start(doc, 14, "Screenshots with Explanation")
    for filename, caption, explanation in SCREENSHOTS:
        image_path = SCREENSHOT_DIR / filename
        if image_path.exists():
            add_figure(doc, image_path, caption, explanation, width=5.95)

    add_chapter_start(doc, 15, "Source Code")
    snippets = [
        (
            "Database initialization and schema definition",
            extract_snippet("def init_db()", "def get_catalog_books"),
            "The `init_db` function is responsible for creating relational tables, backfilling structural changes, seeding default records, and ensuring that the system can start from a known application state. This routine demonstrates how schema management, seed-data bootstrapping, and operational readiness can be coordinated inside a single initialization pipeline.",
        ),
        (
            "Catalog seeding and public metadata import logic",
            extract_snippet("def ensure_catalog", "def lookup_gutendex_text_url"),
            "The catalog population logic is one of the defining features of the project. Instead of relying only on a tiny manual dataset, the application pulls and normalizes public-domain and metadata-driven book records so that the store feels realistic and scalable for demonstration purposes.",
        ),
        (
            "Checkout implementation and order creation",
            extract_snippet('@app.route("/checkout", methods=["GET", "POST"])', '@app.route("/my-orders")'),
            "The checkout route demonstrates core commerce behaviour: cart validation, customer-input capture, order creation, item-line persistence, stock decrement, status initialization, and tracking-number generation. This is a strong indicator that the system goes beyond visual UI work and includes meaningful backend transaction logic.",
        ),
        (
            "Administrative analytics dashboard route",
            extract_snippet('@app.route("/admin/dashboard")', '@app.route("/admin/inventory"'),
            "The admin dashboard route calculates summary statistics and prepares data for KPI cards and reporting panels. It reflects how SQL aggregation and application-layer formatting can be combined to produce manager-friendly analytics without introducing unnecessary complexity.",
        ),
    ]
    for title, code, explanation in snippets:
        add_code_block(doc, title, code, explanation)

    add_chapter_start(doc, 21, "References")
    refs = [
        "[1] Pallets Projects, 'Flask Documentation,' Available: https://flask.palletsprojects.com/. Accessed: 11-May-2026.",
        "[2] SQLite, 'SQLite Documentation,' Available: https://www.sqlite.org/docs.html. Accessed: 11-May-2026.",
        "[3] Open Library, 'APIs,' Available: https://openlibrary.org/developers/api. Accessed: 11-May-2026.",
        "[4] Project Gutenberg, 'About Project Gutenberg,' Available: https://www.gutenberg.org/about/. Accessed: 11-May-2026.",
        "[5] Render, 'Render Documentation,' Available: https://render.com/docs. Accessed: 11-May-2026.",
        "[6] Mozilla Developer Network, 'HTML, CSS and JavaScript Guides,' Available: https://developer.mozilla.org/. Accessed: 11-May-2026.",
    ]
    add_reference_list(doc, refs)
    add_section(
        doc,
        "Technologies and Resources Used in the Project",
        [
            "The implemented Book Store Management System was developed using Python Flask for backend routing and server-side control, SQLite for relational persistence, HTML5 for content structure, CSS3 for layout and styling, and JavaScript for interactivity such as assistant actions and client-side response updates.",
            "The application additionally uses Git and GitHub for source control, Render deployment configuration for live hosting readiness, Open Library metadata for real book records and covers, and Project Gutenberg / Gutendex public-domain text resources for preview-readable content in the browser-based reading module.",
        ],
        bullet_points=[
            "Programming language: Python",
            "Framework: Flask",
            "Database: SQLite",
            "Frontend technologies: HTML5, CSS3, JavaScript",
            "Version control: Git and GitHub",
            "Deployment target: Render",
            "Public metadata sources: Open Library, Project Gutenberg, Gutendex",
        ],
    )
    add_table(
        doc,
        "Table 21.1 Consolidated technology stack used in the project",
        ["Layer", "Technology / Resource", "Use in BOOKVERSE AI"],
        [
            ("Backend", "Python Flask", "Route handling, request processing, session management, and template rendering"),
            ("Database", "SQLite", "Storage of books, reviews, orders, order_items, visitor logs, and support entities"),
            ("Frontend", "HTML5", "Page structure and semantic layout"),
            ("Frontend", "CSS3", "Responsive layout, visual hierarchy, and premium presentation"),
            ("Frontend", "JavaScript", "Assistant interaction, cart actions, and dynamic UI response"),
            ("Version Control", "Git / GitHub", "Source management and deployment collaboration"),
            ("Hosting", "Render configuration", "Live deployment readiness for the Flask application"),
            ("Metadata Source", "Open Library", "Book metadata and cover references"),
            ("Text Preview Source", "Project Gutenberg / Gutendex", "Public-domain reading preview data"),
        ],
    )

    add_chapter_start(doc, 22, "Appendices")
    add_section(
        doc,
        "Appendix A: Sample Route Directory",
        [
            "The route design of the application demonstrates a clear separation between customer-facing URLs, administrative URLs, and API endpoints. This structure improves readability during maintenance and also helps project evaluators verify the flow of responsibility within the backend implementation.",
            "Representative user routes include the home page, store page, book detail page, preview reader, cart, checkout, order tracking, contact, community, and book-club pages. Representative administrative routes include analytics dashboard, inventory management, order supervision, and print-friendly order rendering.",
        ],
    )
    add_table(
        doc,
        "Table 22.1 Selected application routes",
        ["Route", "HTTP Method", "Purpose"],
        [
            ("/", "GET", "Homepage with discovery and recommendation hooks"),
            ("/store", "GET", "Catalog page with search and filters"),
            ("/book/<id>", "GET", "Book detail display"),
            ("/reader/<id>", "GET", "Ten-page reader preview"),
            ("/checkout", "GET/POST", "Checkout screen and order placement"),
            ("/admin/dashboard", "GET", "Administrative KPI dashboard"),
            ("/admin/inventory", "GET/POST", "Inventory management"),
            ("/admin/orders", "GET", "Administrative order viewing"),
            ("/api/store-assistant", "POST", "Assistant-based inventory suggestions"),
        ],
    )
    add_section(
        doc,
        "Appendix B: SQL Schema Notes",
        [
            "The implemented schema is intentionally compact yet representative of a real information system. It separates master data from transactional data, which allows the project to compute stock levels, order totals, and reporting metrics without overloading a single table with unrelated responsibilities.",
            "Additional support tables such as visitor_log and reader_cache demonstrate that the project was designed to support analytics and reading performance, not only storage of bibliographic records.",
        ],
    )
    ddl = extract_snippet("CREATE TABLE IF NOT EXISTS books", '"""\n    )', max_lines=80)
    add_code_block(doc, "Appendix code excerpt: core table definitions", ddl, "This SQL excerpt highlights the schema backbone of the project. It includes the master `books` entity and related tables that capture orders, reviews, wishlist entries, borrowing records, and visitor interactions.")
    add_section(
        doc,
        "Appendix C: User Manual Summary",
        [
            "To use the system as a customer, the user browses the catalog, opens a book, optionally reads the preview, adds the title to the cart, enters personal and shipment information at checkout, and then receives a generated order number and tracking identifier. The same user can later revisit the order-tracking page to inspect progress.",
            "To use the system as an administrator, the authenticated user opens the analytics dashboard, reviews low stock and orders, adds or removes titles through the inventory page, prints order details when required, and uses the metrics cards to monitor operational behaviour. This dual manual demonstrates the multi-role completeness of the application.",
        ],
        bullet_points=[
            "Step 1: Open the homepage and review featured categories and new-arrival sections.",
            "Step 2: Search or filter the store using title, genre, author, language, and rating.",
            "Step 3: Add books to cart and proceed to checkout after authentication.",
            "Step 4: Enter name, phone number, address, city, state, and pincode.",
            "Step 5: Confirm the demo payment and note the generated order and tracking IDs.",
            "Step 6: Log in as admin to inspect dashboard metrics, recent orders, and low-stock alerts.",
        ],
    )
    add_section(
        doc,
        "Appendix D: Observed Live Project Metrics",
        [
            "The final application state used for this report includes measurable runtime values that help demonstrate realism. These values were collected from the current SQLite database rather than being invented for documentation aesthetics.",
            f"At the time of report generation, the catalog contained {stats['books']} titles, {stats['inventory_units']} inventory units, {stats['orders']} orders, {stats['reviews']} reviews, {stats['clubs']} clubs, and {stats['preview_books']} preview-enabled books. The dominant genres were {stats['top_genres']}.",
        ],
    )
    add_page_break(doc)
    add_section(
        doc,
        "Appendix E: Deployment Summary",
        [
            "For live deployment readiness, the project repository was prepared with version control support, dependency declarations, and hosting configuration compatible with Python web-service platforms. The recommended hosting path for the final Flask implementation is Render, because it can execute the backend, templates, and SQLite-backed application without requiring a major architectural rewrite.",
            "Deployment preparation included the creation of a requirements file, GitHub push workflow, and hosted start-command readiness through a production WSGI server. This appendix is relevant in an academic context because it demonstrates that the project was not completed only for offline screenshots; rather, it was structured with the intention of real execution beyond the local development machine.",
            "A deployment-aware project report strengthens evaluation quality because it links software construction with software delivery. By showing how the application can move from code repository to hosted endpoint, the project reflects a fuller software-engineering lifecycle than a report that stops only at implementation and screenshots.",
        ],
        bullet_points=[
            "Repository platform: GitHub",
            "Hosting recommendation: Render web service",
            "Build command: pip install -r requirements.txt",
            "Start command: gunicorn app:app",
        ],
    )
    add_page_break(doc)
    add_section(
        doc,
        "Appendix F: Suggested Demonstration Workflow for Viva or Evaluation",
        [
            "A structured project demonstration improves clarity during viva voce or university evaluation because it allows the examiner to see how the application behaves across customer and administrator roles. The recommended sequence is to begin at the homepage, move into the catalog and filters, open a book detail page, preview the reader, add a book to cart, perform the checkout flow, and finally open the order-tracking page.",
            "After the customer journey is shown, the administrator journey should be demonstrated through the analytics dashboard, low-stock alerts, recent orders, inventory page, and printable order detail screen. This dual-path presentation highlights the central strength of the project: both commerce-facing and management-facing operations are supported in one integrated web platform.",
            "This appendix is intentionally practical. Many project reports document technical design well but do not explain how the finished system should be presented. Including a demonstration workflow makes the report more useful to the student who must defend the project and to the evaluator who wishes to verify the claimed features quickly and systematically.",
        ],
        bullet_points=[
            "Step 1: Show homepage and premium navigation structure.",
            "Step 2: Demonstrate search, filters, and real-cover catalog cards.",
            "Step 3: Open a detail page and launch the sample reader.",
            "Step 4: Add a title to cart, complete checkout, and show generated tracking.",
            "Step 5: Log in as admin and review dashboard, orders, and inventory.",
        ],
    )

    doc.save(DOCX_FILE)


if __name__ == "__main__":
    build_report()
